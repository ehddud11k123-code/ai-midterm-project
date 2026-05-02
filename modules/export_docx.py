from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def _set_font(run, name: str = "맑은 고딕", size: int = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_font(run)


def generate_analysis_docx(text: str, stats: dict, sentiment: dict,
                            readability: dict, keywords: list,
                            summary: list, lang: str = "en",
                            translated: str = None) -> bytes:
    doc = Document()
    _add_heading(doc, "문서 분석 결과", level=0)

    _add_heading(doc, "기본 통계", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "항목"
    hdr[1].text = "값"
    rows = [
        ("단어 수", str(stats.get("word_count", ""))),
        ("문장 수", str(stats.get("sentence_count", ""))),
        ("문단 수", str(stats.get("paragraph_count", ""))),
        ("평균 문장 길이", f"{stats.get('avg_sentence_length', '')} 단어"),
        ("고유 단어 수", str(stats.get("unique_word_count", ""))),
    ]
    for label, value in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
    doc.add_paragraph()

    _add_heading(doc, "핵심 요약", level=1)
    if summary:
        for i, sentence in enumerate(summary, 1):
            _add_paragraph(doc, f"{i}. {sentence}")
    else:
        _add_paragraph(doc, "요약을 생성하기에 텍스트가 너무 짧습니다.")
    doc.add_paragraph()

    _add_heading(doc, "키워드", level=1)
    if keywords:
        kw_text = ", ".join(f"{w}({f})" for w, f in keywords[:15])
        _add_paragraph(doc, kw_text)
    doc.add_paragraph()

    _add_heading(doc, "감정 분석", level=1)
    _add_paragraph(doc, f"판정: {sentiment.get('label', '')}")
    _add_paragraph(doc, f"극성(Polarity): {sentiment.get('polarity', '')}")
    if sentiment.get("reason"):
        _add_paragraph(doc, f"이유: {sentiment['reason']}")
    doc.add_paragraph()

    _add_heading(doc, "가독성", level=1)
    _add_paragraph(doc, f"점수: {readability.get('score', '')}")
    _add_paragraph(doc, f"등급: {readability.get('grade', '')}")

    if translated:
        doc.add_page_break()
        _add_heading(doc, "한국어 번역 (전문)", level=1)
        for para in translated.split("\n"):
            if para.strip():
                _add_paragraph(doc, para.strip())

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_paper_docx(result: dict) -> bytes:
    doc = Document()
    _add_heading(doc, "논문 분석 결과", level=0)

    sections = [
        ("연구 주제", "연구주제"),
        ("주요 기여", "주요기여"),
        ("연구 방법", "연구방법"),
        ("핵심 결과", "핵심결과"),
        ("의의 및 한계", "의의및한계"),
    ]
    for title, key in sections:
        if key in result:
            _add_heading(doc, title, level=1)
            _add_paragraph(doc, result[key])
            doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
